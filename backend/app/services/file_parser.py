"""File parser - extract text from PDF, Word, images, and plain text files."""

import io
import logging
from pathlib import Path

logger = logging.getLogger("seekrefine.parser")


def parse_pdf(content: bytes) -> str:
    """Extract text from PDF using PyMuPDF."""
    import fitz  # pymupdf

    text_parts = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for i, page in enumerate(doc):
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(f"--- Page {i + 1} ---\n{page_text}")

    if not text_parts:
        # PDF might be scanned/image-based, try OCR
        logger.info("PDF has no extractable text, attempting OCR...")
        return _ocr_pdf(content)

    return "\n\n".join(text_parts)


def _ocr_pdf(content: bytes) -> str:
    """OCR a scanned PDF by rendering pages to images."""
    import fitz

    text_parts = []
    with fitz.open(stream=content, filetype="pdf") as doc:
        for i, page in enumerate(doc):
            # Render page to image
            pix = page.get_pixmap(dpi=200)
            img_bytes = pix.tobytes("png")
            page_text = parse_image(img_bytes)
            if page_text.strip():
                text_parts.append(f"--- Page {i + 1} (OCR) ---\n{page_text}")

    return "\n\n".join(text_parts) if text_parts else "(Could not extract text from this PDF)"


def parse_docx(content: bytes) -> str:
    """Extract text from Word .docx file."""
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts) if parts else "(No text found in document)"


def parse_image(content: bytes) -> str:
    """Extract text from image using Tesseract OCR."""
    try:
        import pytesseract
        from PIL import Image

        img = Image.open(io.BytesIO(content))
        text = pytesseract.image_to_string(img, lang="eng+chi_sim")
        return text.strip() if text.strip() else "(No text detected in image)"
    except Exception as e:
        logger.warning(f"OCR failed: {e}")
        return (
            "(OCR is not available. To enable image text extraction, "
            "install Tesseract: https://github.com/tesseract-ocr/tesseract)\n"
            f"Error: {e}"
        )


def parse_file(content: bytes, filename: str) -> str:
    """Parse any supported file type and return extracted text."""
    ext = Path(filename).suffix.lower()
    logger.info(f"Parsing file: {filename} ({len(content)} bytes, type: {ext})")

    # PDF
    if ext == ".pdf":
        return parse_pdf(content)

    # Word
    if ext in (".docx", ".doc"):
        if ext == ".doc":
            return "(Old .doc format is not supported. Please convert to .docx)"
        return parse_docx(content)

    # Images
    if ext in (".png", ".jpg", ".jpeg", ".bmp", ".tiff", ".tif", ".webp"):
        return parse_image(content)

    # Plain text files - try multiple encodings
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb2312", "latin-1"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue

    return "(Unable to read file - unsupported format or encoding)"
